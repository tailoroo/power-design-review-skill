#!/usr/bin/env python3
"""
Word(.docx) 转 Markdown 转换脚本
针对电力设计报告结构优化：标题识别、表格保真、格式保留、噪声过滤

使用方法:
    python3 docx_to_md.py input.docx [-o output.md] [--extract-images]
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def iter_body_elements(doc: DocxDocument):
    """按文档顺序遍历 body 中的段落和表格（不丢内容、不乱序）"""
    body = doc.element.body
    para_idx = 0
    table_idx = 0

    for child in body:
        if child.tag == qn("w:p"):
            # 使用文档自身的段落对象，保留完整上下文
            if para_idx < len(doc.paragraphs):
                yield doc.paragraphs[para_idx]
                para_idx += 1
        elif child.tag == qn("w:tbl"):
            if table_idx < len(doc.tables):
                yield doc.tables[table_idx]
                table_idx += 1


def get_heading_level(para: Paragraph) -> int | None:
    """判断段落标题级别，返回 1-4 或 None"""
    style_name = para.style.name if para.style else ""

    # 样式名匹配
    heading_map = {
        "Heading 1": 1, "Heading 2": 2, "Heading 3": 3, "Heading 4": 4,
        "heading 1": 1, "heading 2": 2, "heading 3": 3, "heading 4": 4,
        "标题 1": 1, "标题 2": 2, "标题 3": 3, "标题 4": 4,
        "标题1": 1, "标题2": 2, "标题3": 3, "标题4": 4,
    }
    if style_name in heading_map:
        return heading_map[style_name]

    # 样式名含 heading + 数字
    match = re.match(r"[Hh]eading\s*(\d)", style_name)
    if match:
        return int(match.group(1))

    # 备选：根据字号和加粗判断
    text = para.text.strip()
    if not text or len(text) > 80:
        return None

    font_size = None
    is_bold = True
    has_runs = False

    for run in para.runs:
        if run.text.strip():
            has_runs = True
            if run.font.size is not None:
                font_size = run.font.size.pt
            if not run.font.bold and not run.bold:
                is_bold = False

    if not has_runs:
        return None

    # 中文报告常用字号：小二(18pt)=章, 三号(16pt)=节, 小三(15pt)=条, 四号(14pt)=款
    if font_size and is_bold:
        if font_size >= 18:
            return 1
        if font_size >= 16:
            return 2
        if font_size >= 14:
            return 3
        if font_size >= 12:
            return 4

    return None


def run_to_md(run) -> str:
    """将单个 run 转为 Markdown 格式文本"""
    text = run.text
    if not text:
        return ""

    # 清理 Word 特殊字符
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 保留有意义的换行，去除尾部换行
    text = text.rstrip("\n")

    bold = run.font.bold or run.bold
    italic = run.font.italic or run.italic

    if bold and italic:
        return f"***{text}***"
    if bold:
        return f"**{text}**"
    if italic:
        return f"*{text}*"
    return text


def paragraph_to_md(para: Paragraph) -> str:
    """将段落转为 Markdown"""
    parts = []
    for run in para.runs:
        parts.append(run_to_md(run))
    text = "".join(parts).strip()

    # 列表检测
    if text:
        text = _convert_list(text)

    return text


def _convert_list(text: str) -> str:
    """将 Word 列表格式转为 Markdown 列表"""
    # 中文数字列表：一、二、三、
    if re.match(r"[一二三四五六七八九十]+、", text):
        return f"- {text}"
    # 数字列表：1. / 1） / 1)
    if re.match(r"\d+[.）)]\s", text):
        return f"- {text}"
    # 字母列表：a) / a） / a.
    if re.match(r"[a-zA-Z][)）.]\s", text):
        return f"- {text}"
    return text


def is_empty_paragraph(para: Paragraph) -> bool:
    """判断是否为空段落"""
    return not para.text.strip()


def is_toc_entry(para: Paragraph) -> bool:
    """判断是否为目录条目"""
    style_name = para.style.name if para.style else ""
    toc_styles = {"TOC Heading", "TOC 1", "TOC 2", "TOC 3", "TOC 4",
                  "toc 1", "toc 2", "toc 3", "toc 4"}
    return style_name in toc_styles


def is_noise_text(text: str) -> bool:
    """判断是否为 Word 转换噪声"""
    noise_patterns = [
        r"错误！未定义书签",
        r"错误！未找到引用源",
        r"错误！不能通过编辑域创建",
        r"REF\s+.*\s+\\h",
    ]
    for pattern in noise_patterns:
        if re.search(pattern, text):
            return True
    return False


def get_cell_text(cell: _Cell) -> str:
    """提取单元格文本，保留 run 格式"""
    parts = []
    for para in cell.paragraphs:
        para_parts = []
        for run in para.runs:
            para_parts.append(run_to_md(run))
        parts.append("".join(para_parts).strip())
    return " ".join(parts).strip()


def table_to_md(table: Table) -> str:
    """将 Word 表格转为 Markdown pipe table"""
    rows = table.rows
    if not rows:
        return ""

    num_cols = max(len(row.cells) for row in rows)

    # 超宽表退化成纯文本块
    if num_cols > 15:
        return _table_to_text_block(table)

    # 构建单元格网格（处理合并单元格）
    grid = _build_cell_grid(table, len(rows), num_cols)

    # 生成 MD 表格
    return _grid_to_md_table(grid, num_cols)


def _build_cell_grid(table: Table, num_rows: int, num_cols: int) -> list[list[str]]:
    """构建二维网格，处理横向和纵向合并"""
    grid: list[list[str | None]] = [[None] * num_cols for _ in range(num_rows)]
    cell_texts: dict[tuple[int, int], str] = {}

    # 第一次遍历：提取每个单元格文本
    for r_idx, row in enumerate(table.rows):
        c_idx = 0
        for cell in row.cells:
            while c_idx < num_cols and grid[r_idx][c_idx] is not None:
                c_idx += 1
            if c_idx >= num_cols:
                break

            text = get_cell_text(cell)
            grid[r_idx][c_idx] = text

            # 检查横向合并（gridSpan）
            tc = cell._tc
            grid_span = tc.find(qn("w:tcPr"))
            if grid_span is not None:
                gs = grid_span.find(qn("w:gridSpan"))
                if gs is not None:
                    span = int(gs.get(qn("w:val"), "1"))
                    for s in range(1, span):
                        if c_idx + s < num_cols:
                            grid[r_idx][c_idx + s] = text

            # 检查纵向合并（vMerge）
            if grid_span is not None:
                vm = grid_span.find(qn("w:vMerge"))
                if vm is not None:
                    merge_val = vm.get(qn("w:val"), "continue")
                    if merge_val == "continue":
                        # 继续合并：从上方行找内容
                        for prev_r in range(r_idx - 1, -1, -1):
                            if grid[prev_r][c_idx] is not None:
                                grid[r_idx][c_idx] = grid[prev_r][c_idx]
                                break
                    # restart 或首次：内容已在 grid 中

            c_idx += 1

    # 填充 None 为空字符串
    result: list[list[str]] = []
    for row in grid:
        result.append([cell if cell is not None else "" for cell in row])
    return result


def _grid_to_md_table(grid: list[list[str]], num_cols: int) -> str:
    """将网格转为 Markdown pipe table"""
    if not grid:
        return ""

    lines = []

    # 表头（第一行）
    header = grid[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * num_cols) + " |")

    # 数据行
    for row in grid[1:]:
        # 补齐列数
        padded = row + [""] * (num_cols - len(row))
        lines.append("| " + " | ".join(padded[:num_cols]) + " |")

    return "\n".join(lines)


def _table_to_text_block(table: Table) -> str:
    """超宽表格退化为纯文本块"""
    lines = []
    for row in table.rows:
        cells = [get_cell_text(cell) for cell in row.cells]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_images(doc: DocxDocument, output_dir: Path) -> list[str]:
    """提取文档中的图片，返回图片路径列表"""
    image_dir = output_dir / "images"
    image_dir.mkdir(exist_ok=True)
    paths = []

    for idx, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            ext = Path(rel.target_ref).suffix or ".png"
            img_path = image_dir / f"image-{idx + 1}{ext}"
            img_path.write_bytes(image_data)
            paths.append(f"images/image-{idx + 1}{ext}")

    return paths


def convert(docx_path: str, output_path: str | None = None, extract_images_flag: bool = False) -> str:
    """主转换函数"""
    doc = Document(docx_path)
    docx_file = Path(docx_path)

    if output_path:
        output_file = Path(output_path)
    else:
        output_file = docx_file.with_suffix(".md")

    output_dir = output_file.parent

    # 提取图片（可选）
    image_paths = []
    if extract_images_flag:
        image_paths = extract_images(doc, output_dir)

    md_lines: list[str] = []
    prev_was_blank = False

    for element in iter_body_elements(doc):
        if isinstance(element, Paragraph):
            para = element

            # 跳过空段落（最多保留一个空行）
            if is_empty_paragraph(para):
                if not prev_was_blank:
                    md_lines.append("")
                    prev_was_blank = True
                continue

            # 跳过目录
            if is_toc_entry(para):
                continue

            # 获取文本
            text = paragraph_to_md(para)
            if not text or is_noise_text(text):
                continue

            # 标题
            level = get_heading_level(para)
            if level:
                prefix = "#" * level
                md_lines.append(f"\n{prefix} {text}\n")
                prev_was_blank = False
                continue

            # 普通段落
            md_lines.append(text)
            prev_was_blank = False

        elif isinstance(element, Table):
            md_table = table_to_md(element)
            if md_table:
                md_lines.append("")
                md_lines.append(md_table)
                md_lines.append("")
                prev_was_blank = True

    # 清理多余空行（最多连续2个）
    result = "\n".join(md_lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    result = result.strip() + "\n"

    output_file.write_text(result, encoding="utf-8")
    print(f"[OK] 转换完成: {output_file}")
    if image_paths:
        print(f"[OK] 提取图片 {len(image_paths)} 张到 {output_dir / 'images'}")
    return str(output_file)


def main():
    parser = argparse.ArgumentParser(description="Word(.docx) 转 Markdown")
    parser.add_argument("input", help="输入 Word 文件路径")
    parser.add_argument("-o", "--output", help="输出 Markdown 文件路径（默认同名.md）")
    parser.add_argument("--extract-images", action="store_true", help="提取图片到 images/ 子目录")
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"[ERROR] 文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    convert(args.input, args.output, args.extract_images)


if __name__ == "__main__":
    main()
