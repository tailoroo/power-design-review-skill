# -*- coding: utf-8 -*-
"""
Markdown转Word文档转换脚本
用于将校审报告Markdown文件转换为Word文档
包含自动检查功能
"""

import re
import sys
import os
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml


def load_format_noise_patterns():
    """加载格式噪音过滤配置"""
    config_path = Path(__file__).parent.parent / "references" / "common-errors.json"
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        return config.get('format_noise_patterns', [])
    except FileNotFoundError:
        print(f"[WARN] 配置文件不存在: {config_path}")
        return []
    except json.JSONDecodeError as e:
        print(f"[WARN] 配置文件JSON格式错误: {e}")
        return []
    except Exception as e:
        print(f"[WARN] 加载配置文件失败: {e}")
        return []


def filter_format_noise(text, patterns):
    """过滤格式噪音

    Args:
        text: 需要过滤的文本
        patterns: 噪音模式列表，每个模式包含 pattern 和 ignore 字段

    Returns:
        过滤后的文本
    """
    if not patterns:
        return text

    for pattern_config in patterns:
        if pattern_config.get('ignore', False):
            try:
                pattern = pattern_config.get('pattern', '')
                if pattern:
                    text = re.sub(pattern, '', text)
            except re.error as e:
                # 正则表达式错误时跳过该模式
                print(f"[WARN] 正则表达式错误，跳过: {pattern_config.get('description', 'unknown')} - {e}")
                continue

    # 清理多余空格
    text = re.sub(r'\s+', ' ', text).strip()
    return text


# 预加载格式噪音模式（模块级别）
_FORMAT_NOISE_PATTERNS = None

def get_format_noise_patterns():
    """获取格式噪音模式（带缓存）"""
    global _FORMAT_NOISE_PATTERNS
    if _FORMAT_NOISE_PATTERNS is None:
        _FORMAT_NOISE_PATTERNS = load_format_noise_patterns()
    return _FORMAT_NOISE_PATTERNS

# 5列表格使用的分类（中文）
VALID_CLASSIFICATIONS = [
    "系统", "电气一次", "电气二次", "土建", "水暖", "概算", "综合", "其他"
]

def setup_document_styles(doc):
    """预先配置文档样式，确保所有样式使用宋体和黑色"""
    from docx.enum.text import WD_COLOR_INDEX
    
    styles = doc.styles
    
    # 设置Heading 1样式 - 大标题，减小字号到16pt
    heading1 = styles['Heading 1']
    heading1.font.name = '宋体'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = None  # 自动颜色（黑色）
    
    # 设置Heading 2样式 - 一级标题，减小字号到14pt
    heading2 = styles['Heading 2']
    heading2.font.name = '宋体'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = None  # 自动颜色（黑色）
    
    # 设置Heading 3样式 - 二级标题，减小字号到12pt
    heading3 = styles['Heading 3']
    heading3.font.name = '宋体'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = None  # 自动颜色（黑色）
    
    # 设置Normal样式 - 正文
    normal = styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(12)
    normal.font.color.rgb = None  # 自动颜色（黑色）
    
    return doc

def set_font_color_black(run):
    """设置字体颜色为黑色"""
    try:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    except:
        pass

def _set_run_font(run, font_name='宋体'):
    """直接设置run的字体 - 使用底层XML"""
    try:
        run.font.name = font_name
        if hasattr(run._element, 'rPr') and run._element.rPr is not None:
            rPr = run._element.rPr
            if rPr.rFonts is None:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                rPr.get_or_add_rFonts()
            rPr.rFonts.set(qn('w:eastAsia'), font_name)
            rPr.rFonts.set(qn('w:hAnsi'), font_name)
    except Exception as e:
        pass

def create_document():
    """创建Word文档"""
    doc = Document()
    setup_document_styles(doc)
    
    # 添加页眉，只显示"设计成品校审单"
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "设计成品校审单"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置页眉字体
        for run in header_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10)
            set_font_color_black(run)
    except Exception as e:
        print(f"添加页眉时出错: {e}")
    
    return doc

def add_heading(doc, text, level=0):
    """添加标题"""
    if level == 0:
        heading = doc.add_heading(text, level=0)
        try:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.bold = True
                _set_run_font(run, '宋体')
                set_font_color_black(run)
        except:
            pass
    elif level == 1:
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.bold = True
            _set_run_font(run, '宋体')
            set_font_color_black(run)
    elif level == 2:
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.bold = True
            _set_run_font(run, '宋体')
            set_font_color_black(run)
    else:
        p = doc.add_paragraph(text)
        p.style = doc.styles['Normal']
        for run in p.runs:
            set_font_color_black(run)
    return doc

def add_paragraph(doc, text, bold=False, italic=False, align=None, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        try:
            p.paragraph_format.first_line_indent = Pt(24)
        except:
            pass
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    _set_run_font(run, '宋体')
    set_font_color_black(run)
    
    if align == 'center':
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    elif align == 'left':
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except:
            pass
    return doc

def parse_markdown_table(md_table):
    """解析Markdown表格"""
    lines = md_table.strip().split('\n')
    tables = []
    current_table = []
    
    for line in lines:
        line = line.strip()
        # 跳过分隔行（包含---的行）
        if '---' in line:
            continue
        if line.startswith('|'):
            content = line.strip('|')
            cells = [cell.strip() for cell in content.split('|')]
            # 过滤掉空的单元格或分隔符单元格
            cells = [cell for cell in cells if cell and not all(c == '-' for c in cell)]
            if cells:  # 只有当有有效单元格时才添加
                current_table.append(cells)
        elif current_table:
            if len(current_table) > 1:
                tables.append(current_table)
            current_table = []
    
    if current_table and len(current_table) > 1:
        tables.append(current_table)
    
    return tables

def set_table_column_widths(table, widths):
    """设置表格列宽（通过XML直接设置，更可靠）"""
    try:
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import parse_xml
        
        tbl = table._tbl
        
        # 禁用表格自动调整
        tblPr = tbl.tblPr
        tblLayout = parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
        tblPr.append(tblLayout)
        
        # 设置每列的宽度
        tblGrid = tbl.tblGrid
        gridCols = tblGrid.gridCol_lst
        for idx, width in enumerate(widths):
            if idx < len(gridCols):
                # 设置列宽（以twips为单位，1英寸=1440 twips）
                width_twips = int(width * 1440)
                gridCols[idx].set(qn('w:w'), str(width_twips))
                gridCols[idx].set(qn('w:type'), 'dxa')
                
        # 设置每个单元格的宽度
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(widths):
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcW = parse_xml(f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{int(widths[col_idx] * 1440)}" w:type="dxa"/>')
                    tcPr.append(tcW)
    except Exception as e:
        print(f"设置列宽时出错: {e}")

def add_markdown_table(doc, md_table):
    """添加Markdown表格到Word"""
    tables = parse_markdown_table(md_table)

    for table_data in tables:
        if not table_data:
            continue

        max_cols = max(len(row) for row in table_data)

        # 检测是否为统计表（包含"分类"和"数量"列的2列表格）
        is_stats_table = False
        if max_cols == 2 and len(table_data) >= 2:
            header = table_data[0]
            if len(header) >= 2 and '分类' in header[0] and '数量' in header[1]:
                is_stats_table = True

        # 检测是否为项目信息表（包含"项目"和"内容"列）
        is_info_table = False
        if max_cols == 2 and len(table_data) >= 2:
            header = table_data[0]
            if len(header) >= 2 and ('项目' in header[0] or '工程名称' in str(table_data)):
                is_info_table = True

        # 项目信息表转为段落形式
        if is_info_table:
            for row_idx, row_data in enumerate(table_data):
                if row_idx == 0:  # 跳过表头
                    continue
                if len(row_data) >= 2:
                    key = clean_markdown_formatting(row_data[0])
                    value = clean_markdown_formatting(row_data[1])
                    if key and value:
                        add_paragraph(doc, f"{key}：{value}")
            return doc

        # 统计表增加占比列
        if is_stats_table:
            # 计算总数
            total = 0
            data_rows = []
            for row_idx, row_data in enumerate(table_data):
                if row_idx == 0:  # 跳过表头
                    continue
                if len(row_data) >= 2:
                    key = clean_markdown_formatting(row_data[0])
                    try:
                        count = int(clean_markdown_formatting(row_data[1]))
                    except:
                        count = 0
                    if key != '合计':
                        total += count
                    data_rows.append((key, count))

            # 创建3列表格（分类、数量、占比）
            table = doc.add_table(rows=len(data_rows) + 1, cols=3)
            table.style = 'Table Grid'
            # 统计表宽度与问题清单一致（总6英寸）
            col_widths = [2.0, 2.0, 2.0]  # 各2英寸，共6英寸
            set_table_column_widths(table, col_widths)

            # 表头
            headers = ['分类', '数量', '占比']
            for col_idx, header_text in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header_text
                try:
                    for run in cell.paragraphs[0].runs:
                        _set_run_font(run, '宋体')
                        set_font_color_black(run)
                        run.font.size = Pt(10)
                        run.bold = True
                except:
                    pass

            # 数据行
            for row_idx, (key, count) in enumerate(data_rows):
                actual_row = row_idx + 1
                cell0 = table.cell(actual_row, 0)
                cell0.text = key
                cell1 = table.cell(actual_row, 1)
                cell1.text = str(count)
                cell2 = table.cell(actual_row, 2)
                if total > 0:
                    percent = count / total * 100
                    cell2.text = f"{percent:.1f}%"
                else:
                    cell2.text = "0.0%"

                for col_idx in range(3):
                    cell = table.cell(actual_row, col_idx)
                    try:
                        for run in cell.paragraphs[0].runs:
                            _set_run_font(run, '宋体')
                            set_font_color_black(run)
                            run.font.size = Pt(10)
                    except:
                        pass

            return doc

        # 普通表格
        table = doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Table Grid'

        # 设置列宽：5列表格专用（序号、分类、问题详细内容、问题位置、修改建议）
        # A4纸可用宽度约6英寸（考虑边距），总列宽不能超过6英寸
        if max_cols == 5:
            # 5列表格：序号、分类、问题详细内容、问题位置、修改建议
            # 总宽度 = 0.5 + 0.8 + 2.4 + 1.2 + 1.1 = 6.0英寸（适配A4页面）
            col_widths = [0.5, 0.8, 2.4, 1.2, 1.1]  # 单位：英寸
            set_table_column_widths(table, col_widths)
        elif max_cols == 2:
            # 2列表格（其他类型）
            col_widths = [2.0, 4.0]
            set_table_column_widths(table, col_widths)

        # 填充表格内容
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                # 清理单元格中的Markdown格式符号
                clean_cell_text = clean_markdown_formatting(cell_text)
                cell.text = clean_cell_text

                try:
                    for run in cell.paragraphs[0].runs:
                        _set_run_font(run, '宋体')
                        set_font_color_black(run)
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                        if row_idx == 0:
                            cell.paragraphs[0].runs[0].bold = True
                except:
                    pass

        return doc

def extract_tables_from_md(md_content):
    """从MD内容中提取所有表格"""
    tables = []
    lines = md_content.split('\n')
    
    in_table = False
    current_table_lines = []
    
    for i, line in enumerate(lines):
        if line.strip().startswith('|') and not in_table:
            in_table = True
            current_table_lines = [line]
            continue
        
        if in_table:
            stripped = line.strip()
            if stripped.startswith('|') or stripped.startswith('|---'):
                current_table_lines.append(line)
            else:
                if len(current_table_lines) > 1:
                    md_table = '\n'.join(current_table_lines)
                    tables.append((i, md_table))
                in_table = False
                current_table_lines = []
    
    if in_table and len(current_table_lines) > 1:
        md_table = '\n'.join(current_table_lines)
        tables.append((len(lines), md_table))
    
    return tables

def check_table_completeness(md_content):
    """检查问题清单完整性"""
    tables = extract_tables_from_md(md_content)
    
    if not tables:
        return False, "NOT_FOUND"
    
    for table_line, table in tables:
        lines = table.split('\n')
        for line in lines:
            if line.strip().startswith('|'):
                cells = [c.strip() for c in line.strip('|').split('|')]
                # 支持5列表格（序号、分类、问题详细内容、问题位置、修改建议）
                # 也支持8列表格（旧格式）
                if len(cells) == 5 or len(cells) == 8:
                    return True, "PASS"
    
    return False, "NOT_5COL"

def check_position_format_from_table(md_content):
    """从表格检查问题位置列是否包含行号格式"""
    tables = extract_tables_from_md(md_content)
    issues = []
    
    for table_line, table in tables:
        lines = table.split('\n')
        header_found = False
        is_valid_table = False
        checked_main = False
        
        for line in lines:
            if re.match(r'^[\|\-\s]+$', line.strip('|').replace('-', '').strip()):
                continue
            if not line.strip().startswith('|'):
                continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            
            # 支持5列表格和8列表格
            if len(cells) == 5 or len(cells) == 8:
                is_valid_table = True
            
            # 只检查第2列(索引1)是"分类"的表格
            if len(cells) >= 2 and cells[1] == '分类':
                header_found = True
                checked_main = True
                continue
            
            # 只检查问题位置列（第4列，索引3）
            # 5列表格：位置在第4列(索引3)
            # 8列表格：位置在第4列(索引3)
            if is_valid_table and header_found and len(cells) >= 4:
                pos_cell = cells[3]
                if pos_cell and re.search(r'第\d+行', pos_cell):
                    issues.append((table_line, "问题位置", pos_cell[:30]))
        
        if checked_main:
            break
    
    if not issues:
        return True, "PASS"
    else:
        issue_msgs = [f"{r[1]}:Ln{r[0]}:{r[2]}" for r in issues[:3]]
        return False, "; ".join(issue_msgs)

def check_classification(md_content):
    """检查问题分类是否有效"""
    tables = extract_tables_from_md(md_content)
    invalid = []
    checked_main = False
    
    for table_line, table in tables:
        if checked_main:
            break
            
        lines = table.split('\n')
        header_found = False
        is_valid_table = False
        
        for line in lines:
            if re.match(r'^[\|\-\s]+$', line.strip('|').replace('-', '').strip()):
                continue
            if not line.strip().startswith('|'):
                continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            
            # 支持5列表格和8列表格
            if len(cells) == 5 or len(cells) == 8:
                is_valid_table = True
            
            if len(cells) >= 2 and cells[1] == '分类':
                header_found = True
                checked_main = True
                continue
            
            if is_valid_table and header_found and len(cells) >= 2:
                cell2 = cells[1]
                # 检查分类是否在有效分类列表中（中文分类）
                if cell2 and cell2 not in VALID_CLASSIFICATIONS:
                    if cell2 not in invalid:
                        invalid.append(cell2[:20])
    
    if not invalid:
        return True, "PASS"
    else:
        return False, f"INVALID:{','.join(invalid[:2])}"

def check_severity(md_content):
    """检查严重程度合理性（5列表格无此列，直接通过）"""
    # 5列表格格式：序号、分类、问题详细内容、问题位置、修改建议
    # 不包含严重程度列，因此直接返回通过
    return True, "PASS"

def check_standard_verification(md_content):
    """检查规范版本问题是否已核实"""
    has_standard = False
    has_verify = False
    
    lines = md_content.split('\n')
    for line in lines:
        line_lower = line.lower()
        if '规范符合性' in line and '分类' not in line:
            if '过期' in line or '更新' in line or 'GB' in line:
                has_standard = True
                if 'codesearch' in line_lower or 'websearch' in line_lower or '核实' in line:
                    has_verify = True
    
    if has_standard and not has_verify:
        return False, "NO_VERIFY"
    return True, "PASS"

def check_filename(md_file):
    """检查文件名规范"""
    filename = os.path.basename(md_file)
    if re.search(r'\d{8}', filename):
        return True, "PASS"
    else:
        return False, "NO_DATE"

def check_review_report(md_file):
    """执行校审报告全面检查"""
    print("=" * 50)
    print("校审报告自动检查")
    print("=" * 50)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    checks = [
        ("文件名规范", check_filename(md_file)),
        ("问题清单完整性", check_table_completeness(md_content)),
        ("问题位置格式", check_position_format_from_table(md_content)),
        ("问题分类一致性", check_classification(md_content)),
        ("规范核实记录", check_standard_verification(md_content)),
    ]
    
    all_passed = True
    
    for check_name, (passed, message) in checks:
        status = "[PASS]" if passed else "[FAIL]"
        print(f"{status} {check_name}")
        if not passed:
            print(f"  -> {message}")
            all_passed = False
    
    print("=" * 50)
    
    return all_passed, checks

def clean_markdown_formatting(text, apply_noise_filter=True):
    """清理Markdown格式符号

    Args:
        text: 需要清理的文本
        apply_noise_filter: 是否应用格式噪音过滤

    Returns:
        清理后的文本
    """
    # 移除标题标记
    text = re.sub(r'^#+\s*', '', text)
    # 处理粗体 **text** -> text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 处理斜体 *text* -> text
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 处理删除线 ~~text~~ -> text
    text = re.sub(r'~~(.+?)~~', r'\1', text)
    # 处理行内代码 `code` -> code
    text = re.sub(r'`(.+?)`', r'\1', text)

    # 应用格式噪音过滤
    if apply_noise_filter:
        patterns = get_format_noise_patterns()
        text = filter_format_noise(text, patterns)

    return text.strip()

def convert_md_to_docx(md_file, docx_file):
    """将Markdown文件转换为Word文档

    Args:
        md_file: 输入的Markdown文件路径
        docx_file: 输出的Word文件路径

    Returns:
        bool: 转换是否成功
    """
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"[ERROR] 文件不存在: {md_file}")
        return False
    except UnicodeDecodeError as e:
        print(f"[ERROR] 文件编码错误: {e}")
        return False
    except Exception as e:
        print(f"[ERROR] 读取文件失败: {e}")
        return False

    try:
        doc = create_document()
    except Exception as e:
        print(f"[ERROR] 无法创建Word文档: {e}")
        return False

    lines = md_content.split('\n')

    # 需要跳过的章节（统计后的冗余内容）
    skip_sections = ['主要问题说明', '校审团队', '输出验证清单']
    current_section = None
    skip_mode = False

    # 用于统计问题数量
    issue_count = 0
    high_priority_keywords = ['不一致', '过期', '错误', '矛盾', '缺失', '漏']

    i = 0
    in_table = False
    current_table_lines = []
    stats_table_processed = False

    while i < len(lines):
        line = lines[i]

        # 跳过分隔线 --- 和空行
        if line.strip() == '---' or line.strip() == '':
            i += 1
            continue

        # 检测二级标题，判断是否需要跳过
        if line.startswith('## ') and not in_table:
            title = clean_markdown_formatting(line)
            # 检查是否是需要跳过的章节
            skip_mode = any(skip_sec in title for skip_sec in skip_sections)
            if not skip_mode:
                current_section = title
                add_heading(doc, title, level=1)
            i += 1
            continue

        # 如果在跳过模式，继续跳过直到下一个二级标题
        if skip_mode:
            if line.startswith('## '):
                # 检查新的章节是否也需要跳过
                title = clean_markdown_formatting(line)
                skip_mode = any(skip_sec in title for skip_sec in skip_sections)
                if not skip_mode:
                    current_section = title
                    add_heading(doc, title, level=1)
            i += 1
            continue

        # 处理一级标题
        if line.startswith('# ') and not in_table:
            title = clean_markdown_formatting(line)
            add_heading(doc, title, level=0)
        # 处理三级标题
        elif line.startswith('### ') and not in_table:
            title = clean_markdown_formatting(line)
            add_heading(doc, title, level=2)
        # 处理粗体文本行
        elif '**' in line and not in_table:
            text = clean_markdown_formatting(line)
            add_paragraph(doc, text, bold=True)
        # 处理表格
        elif line.strip().startswith('|') and not in_table:
            in_table = True
            current_table_lines = [line]
            i += 1
            continue
        elif in_table:
            stripped = line.strip()
            if stripped.startswith('|') or stripped.startswith('|---'):
                current_table_lines.append(line)
            else:
                md_table = '\n'.join(current_table_lines)
                # 检测是否为统计表
                is_stats = '分类' in md_table and '数量' in md_table and '合计' in md_table
                add_markdown_table(doc, md_table)

                # 如果是统计表且还没处理过总结，添加总结章节
                if is_stats and not stats_table_processed:
                    stats_table_processed = True
                    # 解析统计表获取问题数量
                    issue_count = extract_issue_count(md_table)
                    high_count = extract_high_priority_count(md_content, high_priority_keywords)
                    mid_count = issue_count - high_count

                    # 添加总结章节
                    add_heading(doc, "总结", level=1)
                    summary_text = f"本次校审共发现 {issue_count} 个问题，其中："
                    add_paragraph(doc, summary_text)
                    add_paragraph(doc, f"- 高优先级问题：{high_count} 个（数据不一致、规范过期等）")
                    add_paragraph(doc, f"- 中优先级问题：{mid_count} 个（术语错误、描述不完整等）")
                    add_paragraph(doc, "")
                    add_paragraph(doc, "建议设计人员重点核实短路电流计算、规范引用版本等关键问题。")

                in_table = False
                current_table_lines = []
                continue
        # 处理普通段落
        elif not in_table:
            if line.strip():
                clean_text = clean_markdown_formatting(line)
                add_paragraph(doc, clean_text, indent=True)

        i += 1

    try:
        doc.save(docx_file)
        print(f"[OK] Word文档已保存到: {docx_file}")
        return True
    except PermissionError:
        print(f"[ERROR] 无权限写入文件，请检查文件是否被占用: {docx_file}")
        return False
    except Exception as e:
        print(f"[ERROR] 保存Word文档失败: {e}")
        return False


def extract_issue_count(md_table):
    """从统计表中提取问题总数"""
    lines = md_table.strip().split('\n')
    for line in lines:
        if '合计' in line:
            cells = [c.strip() for c in line.strip('|').split('|')]
            for cell in cells:
                try:
                    return int(cell)
                except:
                    continue
    return 0


def extract_high_priority_count(md_content, keywords):
    """从问题清单中统计高优先级问题数量"""
    count = 0
    lines = md_content.split('\n')
    in_issue_table = False

    for line in lines:
        # 检测是否进入问题清单表格
        if '| 序号 |' in line or '|序号|' in line:
            in_issue_table = True
            continue

        if in_issue_table:
            if not line.strip().startswith('|'):
                break
            # 跳过分隔行
            if '---' in line:
                continue
            # 检查问题描述列是否包含高优先级关键词
            cells = [c.strip() for c in line.strip('|').split('|')]
            if len(cells) >= 3:
                description = cells[2]
                if any(kw in description for kw in keywords):
                    count += 1

    return count

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python md_to_word.py <输入MD文件路径>")
        print("示例: python md_to_word.py 校审报告.md")
        print("")
        print("功能说明:")
        print("  1. 自动检查校审报告规范性")
        print("  2. 6项检查全部通过后生成Word文档")
        print("  3. 检查失败时提示修正")
        print("  4. 使用 --force 参数可强制生成（用于模板测试）")
        sys.exit(1)
    
    md_file = sys.argv[1]
    force_mode = '--force' in sys.argv or '-f' in sys.argv
    
    if not os.path.exists(md_file):
        print(f"错误: 文件不存在 - {md_file}")
        sys.exit(1)
    
    if not md_file.endswith('.md'):
        print("错误: 输入文件必须是 .md 格式")
        sys.exit(1)
    
    docx_file = md_file.replace('.md', '.docx')
    
    all_passed, results = check_review_report(md_file)
    
    if not all_passed and not force_mode:
        print("")
        print("检查未通过，请修正后重试。")
        print("常见修正项:")
        print("  - 问题位置改用章节编号，如'第3章第2.1节'、'第4章表4-1'")
        print("  - 分类使用中文名称: 系统、电气一次、电气二次、土建、水暖、概算、综合、其他")
        print("  - 问题清单表格必须是5列: 序号、分类、问题详细内容、问题位置、修改建议")
        print("  - 不记录Word转MD产生的格式问题（如LaTeX公式$220kV$等）")
        print("  - 规范版本问题需调用codesearch/websearch核实")
        sys.exit(1)
    
    if not all_passed and force_mode:
        print("")
        print("[WARN] 检查未通过，但使用强制模式生成文档。")
    
    print("")
    print("正在生成Word文档...")
    print("")
    
    convert_md_to_docx(md_file, docx_file)
    
    print("")
    print("转换完成！")
